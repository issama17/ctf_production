from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Rapport de Projet : Plateforme CTF_LAB', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Module : POO et Programmation Python (Partie II)').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\nDate : 15 Mai 2026').italic = True

    # Section 1
    doc.add_heading('1. Présentation du Projet', level=1)
    doc.add_paragraph(
        "CTF_LAB est une plateforme web interactive de type 'Capture The Flag' (CTF) dédiée à l'apprentissage de la cybersécurité. "
        "Elle permet aux utilisateurs de s'inscrire, de relever des défis techniques (cryptographie, stéganographie, web hacking) "
        "et de progresser dans un classement en temps réel."
    )

    # Section 2
    doc.add_heading('2. Description de l’Application', level=1)
    features = [
        "Gestion des Utilisateurs : Inscription, connexion, profils personnalisés.",
        "Système de Défis : Validation de flags hachés en SHA-256.",
        "Moteur de Score Dynamique : Patron Strategy pour les points dégressifs.",
        "Sécurité : Patron State pour le blocage automatique après 10 échecs.",
        "Scoreboard : Classement dynamique exploitant le polymorphisme.",
        "Labs Vulnérables : Lab SQLi EST Tétouan intégré."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Section 3
    doc.add_heading('3. Modélisation UML', level=1)
    doc.add_paragraph("Note : Les diagrammes UML (Cas d'utilisation et Classes) sont fournis en format Mermaid ci-dessous.")
    
    doc.add_heading('3.1. Diagramme de Cas d’Utilisation (Mermaid)', level=2)
    doc.add_paragraph(
        "mermaid\n"
        "useCaseDiagram\n"
        "    actor \"Participant\" as P\n"
        "    actor \"Administrateur\" as A\n"
        "    P <|-- A\n"
        "    package \"Système CTF_LAB\" {\n"
        "        usecase \"S'inscrire / Se connecter\" as UC1\n"
        "        usecase \"Consulter les défis\" as UC2\n"
        "        usecase \"Soumettre un Flag\" as UC3\n"
        "        usecase \"Télécharger les preuves\" as UC4\n"
        "        usecase \"Consulter le Scoreboard\" as UC5\n"
        "        usecase \"Gérer son profil (Avatar)\" as UC6\n"
        "        usecase \"Accéder au panel Admin\" as UC7\n"
        "        usecase \"Exploiter le Lab SQLi\" as UC8\n"
        "    }\n"
        "    P --> UC1\n"
        "    P --> UC2\n"
        "    P --> UC3\n"
        "    P --> UC4\n"
        "    P --> UC5\n"
        "    P --> UC6\n"
        "    P --> UC8\n"
        "    A --> UC7\n"
    )

    doc.add_heading('3.2. Diagramme de Classes (Architecture POO)', level=2)
    doc.add_paragraph(
        "mermaid\n"
        "classDiagram\n"
        "    class Utilisateur { <<abstract>> -__id : int -__username : str +obtenir_role()* str }\n"
        "    class Participant { +obtenir_role() str }\n"
        "    class Administrateur { +obtenir_role() str }\n"
        "    Utilisateur <|-- Participant\n"
        "    Utilisateur <|-- Administrateur\n"
        "    class Defi { <<abstract>> -__id : str -__points : int +valider_flag() bool }\n"
        "    class DefiWeb { +obtenir_indice() str }\n"
        "    Defi <|-- DefiWeb\n"
        "    class CalculateurScore { <<abstract>> +calculer()* int }\n"
        "    Defi o-- CalculateurScore : utilise\n"
        "    class ContexteDefi { -etat : EtatDefi +essayer_flag() }\n"
        "    class EtatDefi { <<abstract>> +soumettre()* }\n"
        "    ContexteDefi *-- EtatDefi : gère\n"
    )

    # Section 4
    doc.add_heading('4. Explication des Classes Principales', level=1)
    doc.add_paragraph("1. Utilisateur (Encapsulation) : Utilise __ pour le privé.")
    doc.add_paragraph("2. Defi (Héritage) : Base pour les challenges.")
    doc.add_paragraph("3. CalculateurScore (Strategy) : Calcul dégressif.")
    doc.add_paragraph("4. ContexteDefi (State) : Gestion d'état disponible/bloqué.")

    # Section 5
    doc.add_heading('5. Code Source Complet (GitHub)', level=1)
    doc.add_paragraph(
        "Afin de garantir une lecture fluide et de profiter de la coloration syntaxique, "
        "l'intégralité du code source est hébergée sur GitHub. Cliquez sur les liens ci-dessous pour accéder aux fichiers :"
    )

    base_url = "https://github.com/issama17/ctf_platform/blob/main/"
    
    files_to_link = [
        ('Modèles (POO, Encapsulation)', 'models.py'),
        ('Services (Logique métier, Observer)', 'services.py'),
        ('Repositories (Accès BDD)', 'repository.py'),
        ('Application (Façade, Configuration)', 'app.py'),
        ('Routes (Présentation)', 'routes.py'),
        ('Exceptions Personnalisées', 'exceptions.py'),
        ('Point d\'entrée WSGI', 'wsgi.py')
    ]

    for description, filename in files_to_link:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{description} : ").bold = True
        url = f"{base_url}{filename}"
        # Word rend les URLs bleues et cliquables automatiquement
        p.add_run(url).font.color.rgb = RGBColor(0, 0, 255)

    # Save
    doc.save('Rapport_Final_POO_Web.docx')
    print("Document Word avec Hyperliens généré : Rapport_Final_POO_Web.docx")

if __name__ == "__main__":
    create_report()
